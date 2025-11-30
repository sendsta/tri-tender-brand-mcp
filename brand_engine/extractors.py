from __future__ import annotations

import os
from collections import Counter
from typing import List, Tuple

import fitz  # PyMuPDF
import pdfplumber
from docx import Document

from .brand_profile import BrandProfile
from .utils import extract_hex_colors_from_text, image_dominant_colors, safe_filename


def extract_brand_from_files(file_paths: List[str]) -> BrandProfile:
    """
    High-level brand extraction function.

    - Checks PDFs for logos (embedded images), fonts, and hex colors.
    - Checks DOCX for text & hex colors.
    - Checks images directly for dominant colors.
    """
    profile = BrandProfile()

    all_hex_colors = set()
    all_fonts = Counter()
    logo_candidates: List[str] = []
    image_candidates: List[str] = []

    for path in file_paths:
        if not path or not os.path.exists(path):
            continue

        lower = path.lower()
        if lower.endswith(".pdf"):
            pdf_hex, pdf_fonts, pdf_images = _extract_from_pdf(path)
            all_hex_colors.update(pdf_hex)
            all_fonts.update(pdf_fonts)
            logo_candidates.extend(pdf_images)
            image_candidates.extend(pdf_images)
        elif lower.endswith(".docx"):
            docx_hex = _extract_from_docx(path)
            all_hex_colors.update(docx_hex)
        elif lower.endswith((".png", ".jpg", ".jpeg", ".webp")):
            image_candidates.append(path)

    # If we found logo candidates from PDFs, prefer the first
    if logo_candidates:
        profile.logo_path = logo_candidates[0]
    elif image_candidates:
        profile.logo_path = image_candidates[0]

    # Dominant colors: use logo if present, else first image candidate
    dominant_colors: List[str] = []
    if profile.logo_path:
        dominant_colors = image_dominant_colors(profile.logo_path, top_n=5)
    elif image_candidates:
        dominant_colors = image_dominant_colors(image_candidates[0], top_n=5)

    profile.detected_colors = dominant_colors
    profile.hex_colors_in_text = sorted(all_hex_colors)

    # Try to assign primary/secondary/accent from dominant colors or hex in text
    palette_source = dominant_colors or profile.hex_colors_in_text
    if palette_source:
        profile.primary_color = palette_source[0]
    if len(palette_source) > 1:
        profile.secondary_color = palette_source[1]
    if len(palette_source) > 2:
        profile.accent_color = palette_source[2]

    # Fonts
    if all_fonts:
        profile.detected_fonts = [f for f, _ in all_fonts.most_common(5)]
        # Heuristic: if we see "Bold" vs "Regular" we still use the family name
        top_font = profile.detected_fonts[0]
        profile.font_heading = _font_family_from_pdf_name(top_font)
        profile.font_body = _font_family_from_pdf_name(top_font)

    # Ensure chart palette is defined
    profile.ensure_palette()

    return profile


def _extract_from_pdf(path: str) -> Tuple[List[str], Counter, List[str]]:
    """
    Extract hex colors, fonts, and image paths (candidate logos) from a PDF.
    """
    hex_colors = set()
    fonts = Counter()
    image_paths: List[str] = []

    # --- Text & Fonts via pdfplumber ---
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                for c in extract_hex_colors_from_text(text):
                    hex_colors.add(c)

                try:
                    # Some pdfplumber versions expose "chars" with fontname
                    for char in page.chars:
                        fname = char.get("fontname")
                        if fname:
                            fonts[fname] += 1
                except Exception:
                    # Safe fallback if 'chars' not available
                    pass
    except Exception:
        pass

    # --- Images (logos) via PyMuPDF ---
    try:
        doc = fitz.open(path)
        base_dir = os.path.dirname(path) or "."
        logo_out_dir = os.path.join(base_dir, "_extracted_logos")
        os.makedirs(logo_out_dir, exist_ok=True)

        for page_index in range(len(doc)):
            page = doc[page_index]
            img_list = page.get_images(full=True)
            for img_index, img in enumerate(img_list):
                xref = img[0]
                pix = fitz.Pixmap(doc, xref)
                if pix.n > 4:  # handle CMYK
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                out_name = f"pdf_p{page_index}_img{img_index}.png"
                out_path = os.path.join(logo_out_dir, safe_filename(out_name))
                pix.save(out_path)
                image_paths.append(out_path)
                pix = None
    except Exception:
        pass

    return sorted(hex_colors), fonts, image_paths


def _extract_from_docx(path: str) -> List[str]:
    """
    Extract hex colors from DOCX text (e.g. brand guides, CSS snippets).
    """
    hex_colors = set()
    try:
        doc = Document(path)
        for para in doc.paragraphs:
            text = para.text or ""
            for c in extract_hex_colors_from_text(text):
                hex_colors.add(c)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text or ""
                    for c in extract_hex_colors_from_text(text):
                        hex_colors.add(c)
    except Exception:
        pass
    return sorted(hex_colors)


def _font_family_from_pdf_name(font_name: str) -> str:
    """
    Attempt to turn PDF internal font name into a usable CSS font-family.
    """
    if not font_name:
        return "Inter, system-ui, -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif"

    # Remove subset prefixes like "AAAAAA+"
    parts = font_name.split("+", 1)
    fam = parts[-1]
    # Strip style suffixes e.g. "Bold", "Regular"
    fam = fam.replace("Bold", "").replace("Regular", "").strip()
    if not fam:
        fam = "Inter"

    return f"'{fam}', system-ui, -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif"
